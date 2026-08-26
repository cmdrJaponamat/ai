#!/usr/bin/env python3
"""
Приводит DOCX к корпоративному формату АО «Аврора Логистика».

Скрипт не меняет содержание документа. Он:
- задаёт Times New Roman 12 и корпоративный синий для заголовков;
- распознаёт нумерованные разделы и приложения;
- оформляет таблицы, повторяет их шапку и запрещает разрыв строк;
- переносит корпоративный фон/колонтитулы из указанного DOCX-шаблона.

Пример:
  python3 format_corporate_docx.py input.docx \
    --output output.docx \
    --brand-template /path/to/Регламент_учета_ИТ_оборудования_1_4_по_брендбуку.docx
"""
from __future__ import annotations

import argparse
import os
import re
import shutil
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


CORP_BLUE = '004D73'
GRID_COLOR = 'A6A6A6'
REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
ET.register_namespace('', REL_NS)
ET.register_namespace('', CT_NS)
ET.register_namespace('w', W_NS)
ET.register_namespace('r', R_NS)

RE_SECTION_1 = re.compile(r'^\d+\.\s+\S')
RE_SECTION_2 = re.compile(r'^\d+\.\d+\.?\s+\S')
RE_APPENDIX = re.compile(r'^Приложение\s+\d+(?:[.:]|\s)', re.IGNORECASE)
RE_ALL_CAPS_TITLE = re.compile(r'^[А-ЯЁA-Z0-9«»"().,–— -]{12,}$')
# Короткие названия видов внутренних документов не проходят общий порог
# RE_ALL_CAPS_TITLE, но являются первой строкой составного заголовка.
RE_DOCUMENT_KIND_TITLE = re.compile(
    r'^(ПОЛОЖЕНИЕ|РЕГЛАМЕНТ|ИНСТРУКЦИЯ|ПОРЯДОК|НОРМА)$', re.IGNORECASE
)


def set_run_font(run, *, color: str | None = None, bold: bool | None = None):
    run.font.name = 'Times New Roman'
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for key in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        rfonts.set(qn(f'w:{key}'), 'Times New Roman')
    run.font.size = Pt(12)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def clear_cell_shading(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is not None:
        tc_pr.remove(shd)


def set_table_border(table, color: str = GRID_COLOR, size: str = '6'):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = qn(f'w:{edge}')
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f'w:{edge}')
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_row_no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn('w:cantSplit')) is None:
        tr_pr.append(OxmlElement('w:cantSplit'))


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn('w:tblHeader')) is None:
        tr_pr.append(OxmlElement('w:tblHeader'))


def set_style(style, *, color='000000', bold=False, first_line=True, keep=False, before=0, after=5):
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1
    pf.first_line_indent = Cm(1.25) if first_line else Cm(0)
    pf.keep_with_next = keep


def configure_styles(doc):
    set_style(doc.styles['Normal'])
    set_style(doc.styles['Title'], bold=True, first_line=False, keep=True, before=4, after=10)
    set_style(doc.styles['Heading 1'], bold=True, first_line=False, keep=True, before=10, after=4)
    set_style(doc.styles['Heading 2'], bold=True, first_line=False, keep=True, before=7, after=3)
    set_style(doc.styles['Heading 3'], bold=True, first_line=False, keep=True, before=5, after=2)


def classify_paragraph(paragraph, title_seen: bool):
    text = paragraph.text.strip()
    if not text:
        return title_seen
    # В положениях и регламентах «1.1. <длинный текст>» обычно является
    # нумерованным пунктом, а не подзаголовком. Подзаголовком считаем только
    # короткую самостоятельную строку без завершающего знака препинания.
    is_short_subheading = (
        len(text) <= 110 and not text.endswith(('.', ';', ':', '!', '?'))
    )
    if RE_SECTION_2.match(text) and is_short_subheading:
        paragraph.style = 'Heading 2'
    elif RE_SECTION_1.match(text) or RE_APPENDIX.match(text):
        paragraph.style = 'Heading 1'
    elif RE_ALL_CAPS_TITLE.match(text) or RE_DOCUMENT_KIND_TITLE.match(text):
        paragraph.style = 'Title'
        title_seen = True
    return title_seen


def format_paragraphs(doc, auto_headings: bool):
    title_seen = False
    for paragraph in doc.paragraphs:
        if auto_headings:
            title_seen = classify_paragraph(paragraph, title_seen)
        is_heading = paragraph.style.name.startswith('Heading') or paragraph.style.name == 'Title'
        if paragraph.style.name == 'Title':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif paragraph.style.name.startswith('Heading'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.keep_with_next = is_heading
        if not is_heading and paragraph.text.strip():
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.line_spacing = 1
        for run in paragraph.runs:
            set_run_font(run, bold=True if is_heading else None)


def add_table_spacing(table):
    """Adds a 7 pt structural gap after a table, before the following paragraph."""
    following = table._tbl.getnext()
    if following is None or following.tag != qn('w:p'):
        return
    p_pr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '140')
    spacing.set(qn('w:line'), '1')
    spacing.set(qn('w:lineRule'), 'exact')
    p_pr.append(spacing)
    gap = OxmlElement('w:p')
    gap.append(p_pr)
    table._tbl.addnext(gap)


def format_tables(doc):
    for table in reversed(doc.tables):
        table.style = 'Table Grid'
        table.autofit = True
        set_table_border(table)
        if table.rows:
            repeat_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            set_row_no_split(row)
            is_header = row_index == 0
            for cell in row.cells:
                set_cell_margins(cell)
                clear_cell_shading(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.keep_with_next = False
                    for run in paragraph.runs:
                        set_run_font(run, bold=True if is_header else None)
        add_table_spacing(table)


def set_sections(doc, template_path: Path | None):
    template = Document(str(template_path)) if template_path else None
    for index, section in enumerate(doc.sections):
        if template:
            source = template.sections[min(index, len(template.sections) - 1)]
            section.top_margin = source.top_margin
            section.bottom_margin = source.bottom_margin
            section.left_margin = source.left_margin
            section.right_margin = source.right_margin
            section.header_distance = source.header_distance
            section.footer_distance = source.footer_distance
        else:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(2)
            section.right_margin = Cm(1.25)
            section.header_distance = Cm(0.25)
            section.footer_distance = Cm(1.2)


def unique_rel_id(existing):
    numbers = [int(match.group(1)) for value in existing for match in [re.fullmatch(r'rId(\d+)', value)] if match]
    return f'rId{max(numbers, default=0) + 1}'


def add_relationship(root, rel_id, relation_type, target):
    ET.SubElement(root, f'{{{REL_NS}}}Relationship', {'Id': rel_id, 'Type': relation_type, 'Target': target})


def add_override(root, part_name, content_type):
    for node in root.findall(f'{{{CT_NS}}}Override'):
        if node.get('PartName') == part_name:
            return
    ET.SubElement(root, f'{{{CT_NS}}}Override', {'PartName': part_name, 'ContentType': content_type})


def apply_brand_parts(docx_path: Path, template_path: Path):
    """Copies only the corporate headers/footer and their image assets into a DOCX package."""
    source_names = {
        'even': ('word/header1.xml', 'word/_rels/header1.xml.rels'),
        'default': ('word/header2.xml', 'word/_rels/header2.xml.rels'),
        'first': ('word/header3.xml', 'word/_rels/header3.xml.rels'),
        'footer_first': ('word/footer1.xml', 'word/_rels/footer1.xml.rels'),
    }
    output_parts = {
        'even': ('word/corp_header_even.xml', 'word/_rels/corp_header_even.xml.rels'),
        'default': ('word/corp_header_default.xml', 'word/_rels/corp_header_default.xml.rels'),
        'first': ('word/corp_header_first.xml', 'word/_rels/corp_header_first.xml.rels'),
        'footer_first': ('word/corp_footer_first.xml', 'word/_rels/corp_footer_first.xml.rels'),
    }
    image_parts = {
        'word/media/corp_background.jpeg': 'word/media/image2.jpeg',
        'word/media/corp_footer.png': 'word/media/image3.png',
    }

    with zipfile.ZipFile(template_path) as source, zipfile.ZipFile(docx_path) as target:
        payload = {name: target.read(name) for name in target.namelist()}
        for destination, origin in image_parts.items():
            payload[destination] = source.read(origin)
        for key, (origin_xml, origin_rels) in source_names.items():
            destination_xml, destination_rels = output_parts[key]
            payload[destination_xml] = source.read(origin_xml)
            rels = source.read(origin_rels)
            rels = rels.replace(b'media/image2.jpeg', b'media/corp_background.jpeg')
            rels = rels.replace(b'media/image3.png', b'media/corp_footer.png')
            payload[destination_rels] = rels

    rels_root = ET.fromstring(payload['word/_rels/document.xml.rels'])
    existing_ids = [node.get('Id') for node in rels_root.findall(f'{{{REL_NS}}}Relationship')]
    corp_ids = {}
    for key, (_, rel_target) in output_parts.items():
        rel_id = unique_rel_id(existing_ids)
        existing_ids.append(rel_id)
        relation_type = ('http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer'
                         if key.startswith('footer') else
                         'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header')
        target = rel_target.replace('word/_rels/', '').replace('.rels', '')
        add_relationship(rels_root, rel_id, relation_type, target)
        corp_ids[key] = rel_id
    payload['word/_rels/document.xml.rels'] = ET.tostring(rels_root, encoding='utf-8', xml_declaration=True)

    document_root = ET.fromstring(payload['word/document.xml'])
    for sect_pr in document_root.findall(f'.//{{{W_NS}}}sectPr'):
        for node in list(sect_pr):
            if node.tag in (f'{{{W_NS}}}headerReference', f'{{{W_NS}}}footerReference'):
                sect_pr.remove(node)
        refs = [
            (f'{{{W_NS}}}headerReference', 'even', corp_ids['even']),
            (f'{{{W_NS}}}headerReference', 'default', corp_ids['default']),
            (f'{{{W_NS}}}headerReference', 'first', corp_ids['first']),
            (f'{{{W_NS}}}footerReference', 'first', corp_ids['footer_first']),
        ]
        insert_at = 0
        for tag, kind, rel_id in refs:
            node = ET.Element(tag, {f'{{{W_NS}}}type': kind, f'{{{R_NS}}}id': rel_id})
            sect_pr.insert(insert_at, node)
            insert_at += 1
        if sect_pr.find(f'{{{W_NS}}}titlePg') is None:
            sect_pr.insert(insert_at, ET.Element(f'{{{W_NS}}}titlePg'))
    payload['word/document.xml'] = ET.tostring(document_root, encoding='utf-8', xml_declaration=True)

    ct_root = ET.fromstring(payload['[Content_Types].xml'])
    for key, (part, _) in output_parts.items():
        ctype = ('application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml'
                 if key.startswith('footer') else
                 'application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml')
        add_override(ct_root, '/' + part, ctype)
    add_override(ct_root, '/word/media/corp_background.jpeg', 'image/jpeg')
    add_override(ct_root, '/word/media/corp_footer.png', 'image/png')
    payload['[Content_Types].xml'] = ET.tostring(ct_root, encoding='utf-8', xml_declaration=True)

    fd, temporary = tempfile.mkstemp(prefix='corporate-docx-', suffix='.docx', dir=str(docx_path.parent))
    os.close(fd)
    try:
        with zipfile.ZipFile(temporary, 'w', zipfile.ZIP_DEFLATED) as output:
            for name, data in payload.items():
                output.writestr(name, data)
        os.replace(temporary, docx_path)
    except Exception:
        Path(temporary).unlink(missing_ok=True)
        raise


def rebase_on_brand_template(docx_path: Path, template_path: Path):
    """Places the formatted body into the validated corporate template.

    This route is deliberately preferred over low-level package injection: Writer
    and Word preserve the template's page background, header/footer relations and
    page settings as native document parts. Text and tables are copied as XML, so
    direct formatting produced above is retained.
    """
    source = Document(str(docx_path))
    branded = Document(str(template_path))
    target_body = branded.element.body
    target_sect_pr = target_body.sectPr

    for child in list(target_body):
        if child is not target_sect_pr:
            target_body.remove(child)

    insert_at = 0
    for child in source.element.body:
        if child.tag != qn('w:sectPr'):
            copied_child = deepcopy(child)
            # В исходных файлах часто есть промежуточные секции. Их r:id
            # ссылаются на части исходного DOCX и после переноса тела могут
            # совпасть с другим назначением r:id корпоративного шаблона
            # (например, со стилями). Оставляем геометрию секции, но даём ей
            # наследовать фон и колонтитулы шаблона.
            for sect_pr in copied_child.iter(qn('w:sectPr')):
                for reference in list(sect_pr):
                    if reference.tag in (qn('w:headerReference'), qn('w:footerReference')):
                        sect_pr.remove(reference)
            target_body.insert(insert_at, copied_child)
            insert_at += 1

    branded.save(str(docx_path))


def parse_args():
    parser = argparse.ArgumentParser(description='Корпоративное оформление DOCX.')
    parser.add_argument('input', type=Path, help='Исходный DOCX')
    parser.add_argument('--output', type=Path, help='Результирующий DOCX; по умолчанию добавляется суффикс _корпоративный')
    parser.add_argument('--brand-template', type=Path, help='DOCX-шаблон, из которого копируются фон и колонтитулы')
    parser.add_argument('--no-auto-headings', action='store_true', help='Не распознавать нумерованные заголовки и приложения')
    parser.add_argument('--no-brand', action='store_true', help='Не переносить фон и колонтитулы из шаблона')
    return parser.parse_args()


def main():
    args = parse_args()
    if args.input.suffix.lower() != '.docx' or not args.input.is_file():
        raise SystemExit('Нужен существующий файл DOCX.')
    if args.no_brand and args.brand_template:
        raise SystemExit('Выберите либо --brand-template, либо --no-brand.')
    if not args.no_brand and not args.brand_template:
        raise SystemExit('Для корпоративного фона укажите --brand-template; либо явно добавьте --no-brand.')

    output = args.output or args.input.with_name(args.input.stem + '_корпоративный.docx')
    if output.resolve() == args.input.resolve():
        raise SystemExit('Укажите другой --output: исходный файл намеренно не перезаписывается.')
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(args.input, output)

    doc = Document(str(output))
    configure_styles(doc)
    set_sections(doc, args.brand_template)
    format_paragraphs(doc, not args.no_auto_headings)
    format_tables(doc)
    doc.save(output)

    if args.brand_template:
        if not args.brand_template.is_file():
            raise SystemExit(f'Не найден шаблон: {args.brand_template}')
        rebase_on_brand_template(output, args.brand_template)
        # После переноса тела ссылки на стили исходного документа могут не
        # совпадать с идентификаторами стилей корпоративного шаблона. Повторное
        # назначение связывает абзацы с реальными Heading 1/2/Title шаблона.
        branded_doc = Document(str(output))
        configure_styles(branded_doc)
        format_paragraphs(branded_doc, not args.no_auto_headings)
        branded_doc.save(output)

    print(output)


if __name__ == '__main__':
    main()
